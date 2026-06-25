"""出國報告書 docx（對齊行政院出國報告要點附件一）。封面字體 20/26/14 細明體加粗。"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from render.docx_fonts import set_run_font
from render.validators import format_country_region, validate_summary, validate_dates

# 附件二審核表 10 項逐字（renderer contract，不可改字）
REVIEW_ITEMS = [
    "依限繳交出國報告",
    "格式完整（本文具備目的、過程、心得及建議事項）",
    "無抄襲相關資料",
    "內容充實完備",
    "建議具參考價值",
    "送本機關參考或研辦",
    "送上級機關參考",
    "退回補正，原因：(1)不符原核定計畫 (2)以外文撰寫或僅蒐集外文資料 "
    "(3)內容空洞或未涵蓋要項 (4)抄襲 (5)引用未註明來源 (6)電子檔未依格式",
    "公開發表方式：(1)辦理座談會 (2)業務會報提出 (3)其他",
    "其他處理意見及方式",
]
SIGNATURES = ["出國人", "一級單位主管", "機關首長或授權人員"]


def render_review_table_docx(data: dict, out_path: str) -> None:
    doc = Document()
    agency = data.get("agency") or {}
    traveler = data.get("traveler") or {}
    trip = data.get("trip") or {}
    # 表頭標題
    ph = doc.add_paragraph()
    set_run_font(ph.add_run("出國報告審核表"), size_pt=16, bold=True)
    # 附件二官方表頭欄位區塊（報告名稱、出國人、職稱、服務單位、類別、出國期間、繳交日期）
    header_table = doc.add_table(rows=7, cols=2)
    header_table.style = "Table Grid"
    header_rows = [
        ("報告名稱", ""),
        ("出國人姓名", traveler.get("name", "")),
        ("職稱", traveler.get("title", "")),
        ("服務單位", agency.get("full_name", "")),
        ("出國類別", trip.get("purpose_category", "")),
        ("出國期間", f"{trip.get('start_date', '')}～{trip.get('end_date', '')}".strip("～")),
        ("報告繳交日期", ""),
    ]
    for idx, (label, val) in enumerate(header_rows):
        row_cells = header_table.rows[idx].cells
        set_run_font(row_cells[0].paragraphs[0].add_run(label), size_pt=12, bold=True)
        set_run_font(row_cells[1].paragraphs[0].add_run(val), size_pt=12)
    doc.add_paragraph()  # 空行分隔
    # 雙欄勾選表（10 項 + 自我檢核｜主辦審核 兩欄）
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, t in enumerate(["審核項目", "出國人員自我檢核", "計畫主辦機關審核"]):
        set_run_font(hdr[i].paragraphs[0].add_run(t), size_pt=12, bold=True)
    for item in REVIEW_ITEMS:
        row = table.add_row().cells
        set_run_font(row[0].paragraphs[0].add_run(item), size_pt=12)
        set_run_font(row[1].paragraphs[0].add_run("□是 □否"), size_pt=12)
        set_run_font(row[2].paragraphs[0].add_run("□是 □否"), size_pt=12)
    # 三段簽章
    sig = doc.add_table(rows=2, cols=3)
    sig.style = "Table Grid"
    for i, role in enumerate(SIGNATURES):
        set_run_font(sig.rows[0].cells[i].paragraphs[0].add_run(role), size_pt=12, bold=True)
    doc.save(out_path)


def _add_page_number(section):
    fldchar = section.footer.paragraphs[0]
    fldchar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fldchar.add_run()
    fld = run._element.makeelement(qn("w:fldSimple"), {qn("w:instr"): "PAGE"})
    run._element.append(fld)


def render_report_docx(data: dict, out_path: str) -> None:
    agency = data["agency"]
    trip = data["trip"]
    validate_summary(data.get("summary", ""))    # fail fast 在最前
    validate_dates(trip["start_date"], trip["end_date"])  # P2-3：日期邏輯驗證

    doc = Document()
    sec = doc.sections[0]

    # 封面三段（附件一規格）
    p1 = doc.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(p1.add_run(f"出國類別：{trip['purpose_category']}"), size_pt=20, bold=True)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p2.add_run(f"出國報告（出國類別：{trip['purpose_category']}）"), size_pt=26, bold=True)
    # 封面欄位（含單位）
    for label, val in [("服務機關", agency["full_name"]), ("單位", agency.get("unit", "")),
                       ("職稱姓名", f"{data['traveler'].get('title', '')} {data['traveler']['name']}"),
                       ("派赴國家地區", format_country_region(trip)),
                       ("出國期間", f"{trip['start_date']}～{trip['end_date']}")]:
        pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(pc.add_run(f"{label}：{val}"), size_pt=14, bold=True)

    doc.add_page_break()
    # 摘要
    ps = doc.add_paragraph(); set_run_font(ps.add_run("摘要"), size_pt=12, bold=True)
    pst = doc.add_paragraph(); set_run_font(pst.add_run(data["summary"]), size_pt=12)
    # 本文章節
    for head in ("壹、目的", "貳、過程", "參、心得及建議"):
        ph = doc.add_paragraph(); set_run_font(ph.add_run(head), size_pt=12, bold=True)
        pb = doc.add_paragraph(); set_run_font(pb.add_run(""), size_pt=12)

    _add_page_number(sec)
    doc.save(out_path)
