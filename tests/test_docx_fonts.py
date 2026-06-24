"""Task 5: CJK 字型工具測試 — 驗 w:eastAsia 屬性正確寫入 run XML。"""
from docx import Document
from render.docx_fonts import set_run_font


def test_eastasia_attribute_set():
    doc = Document()
    run = doc.add_paragraph().add_run("測試")
    set_run_font(run, name_eastasia="細明體", size_pt=12)
    rpr = run._element.rPr
    rfonts = rpr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
    assert rfonts is not None
    ea = rfonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia")
    assert ea == "細明體"
