"""
Integration test: synthetic example JSON -> four deliverables.

Reads examples/02-sample-agency.*.json, calls all four renderers,
and asserts each output file exists with non-zero size.
"""
import json
import pathlib

import pytest

from render.render_docx import render_report_docx, render_review_table_docx
from render.render_finance_xlsx import render_finance_xlsx
from render.render_html import render_html

_EXAMPLES = pathlib.Path(__file__).parent.parent / "examples"


def _load(name: str) -> dict:
    return json.loads((_EXAMPLES / name).read_text(encoding="utf-8"))


def test_full_pipeline(tmp_path):
    """End-to-end: sample agency data -> report.docx, review.docx, finance.xlsx, pre_trip.html."""
    trip = _load("02-sample-agency.trip.json")
    # Sample summary is 150 chars; validator requires 200-300 Chinese chars.
    trip["summary"] = "本" * 250

    fin = _load("02-sample-agency.trip-finance.json")

    report_path = tmp_path / "report.docx"
    review_path = tmp_path / "review.docx"
    finance_path = tmp_path / "finance.xlsx"
    html_path = tmp_path / "pre_trip.html"

    render_report_docx(trip, str(report_path))
    render_review_table_docx(trip, str(review_path))
    render_finance_xlsx(fin, str(finance_path))
    render_html(trip, str(html_path))

    for f in [report_path, review_path, finance_path, html_path]:
        assert f.exists(), f"{f.name} was not created"
        assert f.stat().st_size > 0, f"{f.name} is empty"


def test_report_docx_contains_key_fields(tmp_path):
    """附件一關鍵欄位：出國人員、期間、目的地、機關."""
    from docx import Document

    trip = _load("02-sample-agency.trip.json")
    trip["summary"] = "本" * 250

    out = tmp_path / "report.docx"
    render_report_docx(trip, str(out))

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Key identifiers present in 附件一 output
    assert trip["traveler"]["name"] in full_text
    assert trip["trip"]["country"] in full_text
    # 「派赴國家地區」欄位必須含城市（附件一「地區」=城市層級）
    assert trip["trip"]["city"] in full_text, f"City '{trip['trip']['city']}' not found in docx cover"


def test_review_table_docx_exists(tmp_path):
    """附件二審核表應產出非空 docx，且第 8 項「退回補正」完整字串存在（renderer contract）."""
    from docx import Document
    from render.render_docx import REVIEW_ITEMS

    trip = _load("02-sample-agency.trip.json")
    trip["summary"] = "本" * 250

    out = tmp_path / "review.docx"
    render_review_table_docx(trip, str(out))

    assert out.exists()
    assert out.stat().st_size > 0

    doc = Document(str(out))
    review_text = "\n".join(p.text for t in doc.tables for row in t.rows for c in row.cells for p in c.paragraphs)
    item8 = REVIEW_ITEMS[7]  # 第 8 項（0-indexed: 7）
    assert item8 in review_text, f"附件二第 8 項字串未出現在 review.docx:\n{item8}"


def test_finance_xlsx_exists(tmp_path):
    """財務計算表應產出非空 xlsx."""
    fin = _load("02-sample-agency.trip-finance.json")

    out = tmp_path / "finance.xlsx"
    render_finance_xlsx(fin, str(out))

    assert out.exists()
    assert out.stat().st_size > 0


def test_html_pre_trip_contains_trip_info(tmp_path):
    """出發前 HTML 應包含行程基本資訊."""
    trip = _load("02-sample-agency.trip.json")
    trip["summary"] = "本" * 250

    out = tmp_path / "pre_trip.html"
    render_html(trip, str(out))

    content = out.read_text(encoding="utf-8")
    assert trip["trip"]["country"] in content
    # 「派赴國家地區」欄位必須含城市
    assert trip["trip"]["city"] in content, f"City '{trip['trip']['city']}' not found in HTML output"
    assert trip["traveler"]["name"] in content
