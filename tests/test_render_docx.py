# tests/test_render_docx.py
import pathlib, pytest
from docx import Document
from render.render_docx import render_report_docx

DATA = {
    "agency": {"full_name": "○○部", "unit": "綜合規劃司"},
    "traveler": {"name": "王測試", "title": "科員"},
    "trip": {"purpose_category": "開會", "country": "日本",
             "start_date": "2027-03-01", "end_date": "2027-03-05"},
    "summary": "本" * 250,
}

def test_renders_required_sections(tmp_path):
    out = tmp_path / "report.docx"
    render_report_docx(DATA, str(out))
    doc = Document(str(out))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "壹、目的" in texts
    assert "貳、過程" in texts
    assert "參、心得及建議" in texts
    assert "○○部" in texts          # 機關名從 agency 讀
    assert "綜合規劃司" in texts      # 單位欄

def test_bad_summary_fails(tmp_path):
    bad = {**DATA, "summary": "太短"}
    with pytest.raises(Exception):
        render_report_docx(bad, str(tmp_path / "x.docx"))


def test_body_sections_have_grey_writing_hints(tmp_path):
    """本文三章節須有灰色撰寫提示（提醒充實本文、非留空），且為灰色（明示待刪）。"""
    from docx.shared import RGBColor
    out = tmp_path / "report.docx"
    render_report_docx(DATA, str(out))
    doc = Document(str(out))
    hint_paras = [p for p in doc.paragraphs if "撰寫提示" in p.text]
    assert len(hint_paras) == 3, f"應有 3 段撰寫提示，實得 {len(hint_paras)}"
    # 過程章節的提示須提及逐字稿/筆記（引導使用者用素材充實本文）
    assert any("逐字稿" in p.text and "機密" in p.text for p in hint_paras)
    # 提示須為灰色
    for p in hint_paras:
        assert p.runs and p.runs[0].font.color.rgb == RGBColor(0x80, 0x80, 0x80)
