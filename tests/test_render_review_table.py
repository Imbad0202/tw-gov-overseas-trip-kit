# tests/test_render_review_table.py
from docx import Document
from render.render_docx import render_review_table_docx, REVIEW_ITEMS

DATA = {"agency": {"full_name": "○○部", "head_title": "部長"},
        "traveler": {"name": "王測試", "title": "科員"},
        "trip": {"purpose_category": "開會",
                 "start_date": "2027-09-01", "end_date": "2027-09-05"}}


def test_ten_items_verbatim(tmp_path):
    out = tmp_path / "review.docx"
    render_review_table_docx(DATA, str(out))
    doc = Document(str(out))
    full = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert len(REVIEW_ITEMS) == 10, f"審核表應為 10 項，實際 {len(REVIEW_ITEMS)}"
    for item in REVIEW_ITEMS:
        assert item in full, f"審核表缺項或後半被改：『{item}』"


def test_three_signature_blocks(tmp_path):
    out = tmp_path / "review.docx"
    render_review_table_docx(DATA, str(out))
    doc = Document(str(out))
    full = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "出國人" in full and "一級單位主管" in full and "機關首長" in full


def test_header_fields_present(tmp_path):
    """附件二表頭欄位：出國人姓名、職稱、服務單位、出國類別、出國期間、報告繳交日期。"""
    out = tmp_path / "review_header.docx"
    render_review_table_docx(DATA, str(out))
    doc = Document(str(out))
    # 取全部段落文字 + 表格文字
    para_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    full = para_text + "\n" + table_text
    for label in ("出國人姓名", "職稱", "服務單位", "出國類別", "出國期間", "報告繳交日期"):
        assert label in full, f"附件二表頭缺欄位：『{label}』"
    # 確認 traveler/trip/agency 資料已帶入
    assert "王測試" in full
    assert "科員" in full
    assert "○○部" in full
    assert "開會" in full
    assert "2027-09-01" in full


def test_header_missing_fields_no_error(tmp_path):
    """缺少 traveler/trip/agency 欄位時不報錯，以空白 guard 輸出。"""
    minimal_data = {}
    out = tmp_path / "review_minimal.docx"
    render_review_table_docx(minimal_data, str(out))  # 不應拋出任何 exception
    doc = Document(str(out))
    assert doc is not None
